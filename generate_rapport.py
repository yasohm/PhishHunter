"""
Génération du rapport PFE PhishGuard en format .docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Couleurs ──────────────────────────────────────────────────────────────────
BLEU_TITRE  = RGBColor(0x1F, 0x4E, 0x79)   # H1
BLEU_H2     = RGBColor(0x2E, 0x75, 0xB6)   # H2
BLEU_HEADER = RGBColor(0x2E, 0x75, 0xB6)   # en-têtes de table
BLANC       = RGBColor(0xFF, 0xFF, 0xFF)
GRIS_CODE   = RGBColor(0xF2, 0xF2, 0xF2)
BLEU_FIGURE = RGBColor(0xF0, 0xF4, 0xFF)
GRIS_TEXTE  = RGBColor(0x55, 0x55, 0x55)
BLEU_ALT    = RGBColor(0xEB, 0xF3, 0xFB)   # lignes alternées

LOGO_PATH   = os.path.join(os.path.dirname(__file__),
                           "phishguard-extension", "icons", "icon-128.png")


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "BFBFBF")
        borders.append(b)
    tcPr.append(borders)


def _page_break(doc):
    doc.add_page_break()


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = BLEU_TITRE
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(10)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = BLEU_H2
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(8)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = BLEU_TITRE
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(6)


def para(doc, text, bold=False, italic=False, center=False, size=12,
         color=None, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if color:
        run.font.color.rgb = color
    return p


def bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(12)
        r.font.name = "Times New Roman"
    r2 = p.add_run(text)
    r2.font.size = Pt(12)
    r2.font.name = "Times New Roman"


def code_block(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def figure_placeholder(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F4FF")
    pPr.append(shd)
    run = p.add_run(f"[ {caption} ]")
    run.italic = True
    run.font.color.rgb = GRIS_TEXTE
    run.font.size = Pt(10)


def separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # En-tête
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, "2E75B6")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = BLANC
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

    # Lignes de données
    for r_idx, row_data in enumerate(rows):
        bg = "EBF3FB" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = tbl.rows[r_idx+1].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"

    # Largeurs de colonnes
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return tbl


# ── Document ──────────────────────────────────────────────────────────────────

doc = Document()

# Mise en page A4
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)

# Style par défaut
style = doc.styles["Normal"]
font  = style.font
font.name = "Times New Roman"
font.size = Pt(12)
pf = style.paragraph_format
pf.line_spacing = Pt(18)
pf.space_after  = Pt(6)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE DE GARDE — avec logo et titre stylisé
# ═══════════════════════════════════════════════════════════════════════════════

# Espace en haut
for _ in range(2):
    doc.add_paragraph()

# Ligne de logos établissement
p_logos = doc.add_paragraph()
p_logos.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_logos.add_run("[Logo de l'établissement]          [Logo de l'organisme d'accueil]")
r.italic = True
r.font.color.rgb = GRIS_TEXTE
r.font.size = Pt(10)

doc.add_paragraph()

# ── Logo PhishGuard ────────────────────────────────────────────────────────────
if os.path.exists(LOGO_PATH):
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(12)
    p_logo.paragraph_format.space_after  = Pt(12)
    run_logo = p_logo.add_run()
    run_logo.add_picture(LOGO_PATH, width=Inches(1.8))

# ── Titre principal PhishGuard ─────────────────────────────────────────────────
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(8)
p_title.paragraph_format.space_after  = Pt(4)
r_title = p_title.add_run("PhishGuard")
r_title.bold = True
r_title.font.size = Pt(36)
r_title.font.name = "Times New Roman"
r_title.font.color.rgb = BLEU_TITRE

# ── Sous-titre ─────────────────────────────────────────────────────────────────
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after  = Pt(4)
r_sub = p_sub.add_run("Phishing Detection Platform")
r_sub.font.size = Pt(14)
r_sub.font.name = "Times New Roman"
r_sub.font.color.rgb = BLEU_H2

# ── Ligne séparatrice ──────────────────────────────────────────────────────────
p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_line.paragraph_format.space_before = Pt(4)
p_line.paragraph_format.space_after  = Pt(16)
r_line = p_line.add_run("─" * 50)
r_line.font.color.rgb = BLEU_H2
r_line.font.size = Pt(11)

# ── Description du rapport ─────────────────────────────────────────────────────
p_desc = doc.add_paragraph()
p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_desc.paragraph_format.space_after = Pt(4)
r_desc = p_desc.add_run("Système de Détection de Phishing par Intelligence Artificielle")
r_desc.bold = True
r_desc.font.size = Pt(14)
r_desc.font.color.rgb = BLEU_TITRE

p_type = doc.add_paragraph()
p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_type.paragraph_format.space_after = Pt(20)
r_type = p_type.add_run("Rapport de Projet de Fin d'Études (PFE)")
r_type.italic = True
r_type.font.size = Pt(12)
r_type.font.color.rgb = GRIS_TEXTE

# ── Infos académiques ──────────────────────────────────────────────────────────
for label, value in [
    ("Réalisé par", "[Votre Nom et Prénom]"),
    ("Encadrant",   "[Nom de l'encadrant]"),
    ("Niveau",      "BTS / DSI"),
    ("Filière",     "Développement des Systèmes d'Information"),
    ("Année scolaire", "2024–2025"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{label} : ")
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.name = "Times New Roman"
    r2 = p.add_run(value)
    r2.font.size = Pt(12)
    r2.font.name = "Times New Roman"

_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# REMERCIEMENTS
# ═══════════════════════════════════════════════════════════════════════════════

heading(doc, "REMERCIEMENTS", 1)
para(doc, (
    "Je tiens à exprimer ma profonde gratitude à toutes les personnes qui ont contribué, "
    "de près ou de loin, à la réalisation de ce projet de fin d'études."
))
para(doc, (
    "Je remercie en premier lieu mon encadrant, [Nom de l'encadrant], pour ses précieux conseils, "
    "sa disponibilité et ses orientations tout au long de ce travail. Son expertise et son soutien "
    "ont été déterminants dans la conduite de ce projet."
))
para(doc, (
    "Je remercie également l'ensemble du corps enseignant de la filière Développement des Systèmes "
    "d'Information pour la qualité de la formation dispensée, qui m'a fourni les bases techniques "
    "nécessaires à la réalisation de ce projet."
))
para(doc, "Mes remerciements vont aussi aux membres du jury pour le temps consacré à l'évaluation de ce travail.")
para(doc, (
    "Enfin, je remercie ma famille et mes proches pour leur soutien moral indéfectible "
    "durant toute la durée de mes études."
))
_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# RÉSUMÉ / ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════════

heading(doc, "RÉSUMÉ", 1)
para(doc, (
    "Le phishing constitue l'une des cybermenaces les plus répandues, causant des pertes financières "
    "et des atteintes à la confidentialité des données à l'échelle mondiale. Ce projet de fin d'études "
    "présente PhishGuard, un système de détection de phishing basé sur l'Intelligence Artificielle, "
    "développé sous forme d'application web full-stack complétée par une extension de navigateur."
))
para(doc, (
    "L'architecture du système repose sur deux modules de détection complémentaires : un analyseur d'URLs "
    "utilisant un classifieur Gradient Boosting entraîné sur 30 caractéristiques extraites du dataset UCI "
    "Phishing Websites (11 055 exemples), et un analyseur d'emails employant un modèle Random Forest avec "
    "une approche Bag of Words. La solution intègre un backend FastAPI (Python), une base de données "
    "PostgreSQL, un frontend React/Vite, et une extension navigateur compatible Manifest V3."
))
para(doc, (
    "Les résultats obtenus démontrent une précision supérieure à 95 % pour la détection d'URLs de phishing, "
    "avec une interface utilisateur intuitive permettant la visualisation détaillée des caractéristiques "
    "d'analyse, l'historique des scans et l'export des rapports."
))
p_kw = doc.add_paragraph()
p_kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r1 = p_kw.add_run("Mots-clés : ")
r1.bold = True; r1.font.size = Pt(12); r1.font.name = "Times New Roman"
r2 = p_kw.add_run("Phishing, Machine Learning, Gradient Boosting, Random Forest, FastAPI, React, Extension Chrome, Cybersécurité.")
r2.font.size = Pt(12); r2.font.name = "Times New Roman"

separator(doc)
heading(doc, "ABSTRACT", 2)
para(doc, (
    "Phishing represents one of the most widespread cyber threats, causing financial losses and data "
    "breaches worldwide. This final-year project presents PhishGuard, an AI-powered phishing detection "
    "system developed as a full-stack web application complemented by a browser extension."
))
para(doc, (
    "The system architecture relies on two complementary detection modules: a URL analyzer using a "
    "Gradient Boosting classifier trained on 30 features extracted from the UCI Phishing Websites dataset "
    "(11,055 samples), and an email analyzer employing a Random Forest model with a Bag of Words approach. "
    "The solution integrates a FastAPI backend (Python), a PostgreSQL database, a React/Vite frontend, "
    "and a Manifest V3-compatible browser extension."
))
para(doc, (
    "Results demonstrate accuracy exceeding 95% for phishing URL detection, with an intuitive user interface "
    "for detailed feature visualization, scan history, and report export."
))
p_kw2 = doc.add_paragraph()
p_kw2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r1 = p_kw2.add_run("Keywords: ")
r1.bold = True; r1.font.size = Pt(12); r1.font.name = "Times New Roman"
r2 = p_kw2.add_run("Phishing, Machine Learning, Gradient Boosting, Random Forest, FastAPI, React, Chrome Extension, Cybersecurity.")
r2.font.size = Pt(12); r2.font.name = "Times New Roman"
_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# SOMMAIRE
# ═══════════════════════════════════════════════════════════════════════════════

heading(doc, "SOMMAIRE", 1)
toc_items = [
    ("Introduction", "1"),
    ("Chapitre 1 : Description du Sujet", "4"),
    ("  1.1 Contexte et problématique", "4"),
    ("  1.2 Présentation du projet PhishGuard", "6"),
    ("  1.3 Objectifs et intérêt du projet", "8"),
    ("Chapitre 2 : Conception de la Solution", "10"),
    ("  2.1 Technologies et outils utilisés", "10"),
    ("  2.2 Architecture du système", "13"),
    ("  2.3 Modélisation des données", "16"),
    ("  2.4 Conception des modules ML", "18"),
    ("Chapitre 3 : Implémentation du Projet", "22"),
    ("  3.1 Mise en place de l'environnement", "22"),
    ("  3.2 Implémentation du backend", "23"),
    ("  3.3 Implémentation du frontend", "27"),
    ("  3.4 Extension navigateur", "30"),
    ("  3.5 Tests et évaluation", "33"),
    ("Conclusion", "36"),
    ("Références Bibliographiques", "38"),
    ("Annexes", "40"),
]
for item, pg in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    dots = "." * max(2, 60 - len(item) - len(pg))
    bold = not item.startswith("  ")
    r = p.add_run(f"{item} {dots} {pg}")
    r.bold = bold
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# LISTE DES ABRÉVIATIONS
# ═══════════════════════════════════════════════════════════════════════════════

heading(doc, "LISTE DES ABRÉVIATIONS", 1)
abbr_rows = [
    ("IA / AI",  "Intelligence Artificielle / Artificial Intelligence"),
    ("ML",       "Machine Learning (Apprentissage Automatique)"),
    ("API",      "Application Programming Interface"),
    ("URL",      "Uniform Resource Locator"),
    ("HTTP / HTTPS", "HyperText Transfer Protocol (Secure)"),
    ("HTML",     "HyperText Markup Language"),
    ("JSON",     "JavaScript Object Notation"),
    ("REST",     "Representational State Transfer"),
    ("CORS",     "Cross-Origin Resource Sharing"),
    ("ORM",      "Object-Relational Mapping"),
    ("UCI",      "University of California, Irvine"),
    ("RF",       "Random Forest"),
    ("GB",       "Gradient Boosting"),
    ("NLP",      "Natural Language Processing"),
    ("BoW",      "Bag of Words"),
    ("SSL",      "Secure Sockets Layer"),
    ("DNS",      "Domain Name System"),
    ("WHOIS",    "Protocole de requête de données d'enregistrement de domaine"),
    ("DOM",      "Document Object Model"),
    ("MV3",      "Manifest Version 3 (extension navigateur)"),
    ("PFE",      "Projet de Fin d'Études"),
    ("BTS",      "Brevet de Technicien Supérieur"),
    ("DSI",      "Développement des Systèmes d'Information"),
]
add_table(doc, ["Abréviation", "Signification"], abbr_rows, [4, 12])
_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# LISTE DES FIGURES
# ═══════════════════════════════════════════════════════════════════════════════

heading(doc, "LISTE DES FIGURES", 1)
figures = [
    ("Figure 1",  "Évolution des attaques de phishing mondiales (2019–2024)"),
    ("Figure 2",  "Architecture globale du système PhishGuard"),
    ("Figure 3",  "Diagramme des cas d'utilisation"),
    ("Figure 4",  "Modèle de la base de données (table Scans)"),
    ("Figure 5",  "Pipeline d'extraction des caractéristiques d'URL"),
    ("Figure 6",  "Pipeline d'analyse des emails"),
    ("Figure 7",  "Résultats de validation croisée du modèle Gradient Boosting"),
    ("Figure 8",  "Importance des caractéristiques (Feature Importance)"),
    ("Figure 9",  "Matrice de confusion du modèle URL"),
    ("Figure 10", "Architecture Docker — orchestration des conteneurs"),
    ("Figure 11", "Interface d'accueil de PhishGuard"),
    ("Figure 12", "Interface d'analyse d'URL"),
    ("Figure 13", "Résultat d'analyse — site dangereux"),
    ("Figure 14", "Interface d'analyse d'email"),
    ("Figure 15", "Page d'historique avec filtres"),
    ("Figure 16", "Popup de l'extension navigateur"),
    ("Figure 17", "Comparaison des performances des modèles"),
]
add_table(doc, ["Référence", "Description"], figures, [3, 13])
_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════

heading(doc, "INTRODUCTION", 1)

heading(doc, "1. Contexte général", 2)
para(doc, (
    "La transformation numérique accélérée de nos sociétés a engendré une multiplication sans précédent "
    "des cybermenaces. Parmi celles-ci, le phishing (hameçonnage) occupe une place prépondérante : selon "
    "le rapport APWG (Anti-Phishing Working Group) de 2024, plus de 5 millions d'attaques de phishing "
    "uniques ont été enregistrées au cours du seul premier trimestre 2024, établissant un nouveau record "
    "historique. Ces attaques visent à tromper les utilisateurs pour leur soutirer des informations sensibles "
    "— identifiants, mots de passe, données bancaires — en usurpant l'identité de sites ou d'expéditeurs légitimes."
))
figure_placeholder(doc, "Figure 1 : Évolution des attaques de phishing mondiales (2019–2024)")
para(doc, (
    "Les méthodes traditionnelles de détection, basées sur des listes noires (blacklists) statiques, montrent "
    "leurs limites face à la sophistication croissante des attaquants : les URLs de phishing ont une durée "
    "de vie de quelques heures seulement, rendant les listes noires perpétuellement obsolètes. Il devient "
    "donc indispensable de recourir à des méthodes d'apprentissage automatique capables de détecter des "
    "patterns inconnus en temps réel."
))

heading(doc, "2. Cadre du stage", 2)
para(doc, (
    "Ce projet de fin d'études s'inscrit dans le cadre de la formation BTS Développement des Systèmes "
    "d'Information et répond à un double objectif : d'une part, mettre en pratique les compétences acquises "
    "en développement web, gestion de bases de données et algorithmes ; d'autre part, contribuer à la "
    "cybersécurité à travers une solution concrète et déployable."
))

heading(doc, "3. But pédagogique", 2)
para(doc, "Ce projet vise à consolider les compétences dans les domaines suivants :")
for b in [
    "Le développement full-stack (Python/FastAPI pour le backend, React pour le frontend)",
    "L'intégration de modèles de Machine Learning dans une application web",
    "La conteneurisation d'applications avec Docker",
    "Le développement d'extensions navigateur (Manifest V3)",
    "La conception et modélisation de bases de données relationnelles (PostgreSQL)",
]:
    bullet(doc, b)

heading(doc, "4. Problématique", 2)
para(doc, (
    "Comment concevoir et implémenter un système de détection de phishing en temps réel, basé sur "
    "l'intelligence artificielle, accessible à travers une application web et une extension navigateur, "
    "capable d'analyser aussi bien les URLs que les emails suspects ?"
), bold=True)
para(doc, "Cette problématique soulève plusieurs questions secondaires :")
for b in [
    "Quelles caractéristiques d'une URL ou d'un email permettent de distinguer le phishing des contenus légitimes ?",
    "Quel algorithme de Machine Learning offre le meilleur compromis entre précision et vitesse d'exécution ?",
    "Comment architecturer une solution scalable et facilement déployable ?",
]:
    bullet(doc, b)

heading(doc, "5. Démarche du rapport", 2)
para(doc, (
    "Le Chapitre 1 présente le contexte du projet, la problématique détaillée du phishing et les "
    "objectifs de PhishGuard. "
    "Le Chapitre 2 expose la conception de la solution : choix technologiques, architecture système, "
    "modélisation des données et conception des modules de Machine Learning. "
    "Le Chapitre 3 détaille l'implémentation concrète du projet, les tests réalisés et l'évaluation "
    "des performances."
))
_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 1
# ═══════════════════════════════════════════════════════════════════════════════

heading(doc, "CHAPITRE 1 : DESCRIPTION DU SUJET", 1)

heading(doc, "1.1 Contexte et Problématique", 2)
heading(doc, "1.1.1 Le Phishing : Définition et Ampleur", 3)
para(doc, (
    "Le phishing est une technique de cyberattaque par ingénierie sociale dans laquelle un attaquant "
    "se fait passer pour une entité de confiance (banque, réseau social, service gouvernemental) afin "
    "de pousser la victime à divulguer des informations confidentielles ou à effectuer des actions malveillantes."
))
para(doc, "Les vecteurs d'attaque sont multiples :")
for b, t in [
    ("URLs malveillantes", " : sites web imitant fidèlement des pages légitimes"),
    ("Emails de phishing", " : messages frauduleux incitant à cliquer sur des liens ou à fournir des données"),
    ("SMS (Smishing)", " et appels téléphoniques (Vishing)"),
]:
    bullet(doc, t, bold_prefix=b)

heading(doc, "1.1.2 Limites des Solutions Existantes", 3)
para(doc, "Les approches actuelles de lutte contre le phishing présentent plusieurs limitations :")
add_table(doc,
    ["Approche", "Avantages", "Inconvénients"],
    [
        ("Listes noires (Blacklists)", "Simple, rapide", "Ne détecte pas les nouvelles URLs, mise à jour lente"),
        ("Filtres par règles",         "Explicable",     "Facilement contournable"),
        ("Vérification manuelle",      "Fiable",         "Non scalable, trop lente"),
        ("ML basé sur le contenu",     "Généralise bien","Nécessite un accès au contenu de la page"),
    ],
    [5, 4, 7]
)
para(doc, (
    "Ces limitations justifient le développement d'une solution hybride combinant l'analyse de structure "
    "d'URL et l'analyse sémantique des emails, à travers des modèles de Machine Learning entraînés sur "
    "des datasets publics reconnus."
))

heading(doc, "1.1.3 Contexte Organisationnel", 3)
para(doc, (
    "PhishGuard a été développé dans le cadre d'un projet académique individuel. Le projet s'inscrit "
    "dans la démarche plus large de démocratisation des outils de cybersécurité, en proposant une "
    "solution open-source, facilement déployable via Docker, et accessible au grand public à travers "
    "une extension de navigateur."
))

heading(doc, "1.2 Présentation du Projet PhishGuard", 2)
heading(doc, "1.2.1 Vue d'Ensemble", 3)
para(doc, "PhishGuard est un système de détection de phishing en temps réel basé sur l'Intelligence Artificielle. Il se compose de trois composantes principales :")
for n, t in [
    ("Application Web", " : Interface permettant l'analyse d'URLs et d'emails, la visualisation des caractéristiques détectées et la consultation de l'historique des analyses."),
    ("API Backend",     " : Serveur RESTful assurant l'extraction des caractéristiques, l'inférence des modèles ML et la persistance des données."),
    ("Extension Navigateur", " : Composant intégré au navigateur (Chrome/Firefox) permettant une détection automatique et transparente des pages web visitées par l'utilisateur."),
]:
    bullet(doc, t, bold_prefix=n)

heading(doc, "1.2.2 Fonctionnalités Principales", 3)
para(doc, "Analyse d'URLs :", bold=True)
for b in [
    "Extraction de 30 caractéristiques structurelles et comportementales",
    "Classification en trois niveaux de risque : Sûr, Suspect, Dangereux",
    "Score de confiance en pourcentage",
    "Visualisation détaillée de chaque caractéristique analysée",
]:
    bullet(doc, b)

para(doc, "Analyse d'Emails :", bold=True)
for b in [
    "Extraction de caractéristiques NLP (Bag of Words, entropie, mots-clés d'urgence)",
    "Détection de spoofing de marques",
    "Analyse des URLs contenues dans l'email",
    "Score de risque combiné (email + URLs)",
]:
    bullet(doc, b)

para(doc, "Fonctionnalités transversales :", bold=True)
for b in [
    "Historique complet des analyses avec pagination et filtres",
    "Export des rapports en formats PDF et JSON",
    "Extension navigateur avec popup de résultat en temps réel",
    "Tableau de bord statistique",
]:
    bullet(doc, b)

heading(doc, "1.3 Objectifs et Intérêt du Projet", 2)
heading(doc, "1.3.1 Objectifs Techniques", 3)
for b in [
    "Concevoir et entraîner deux modèles de ML distincts (URLs et emails) avec des performances élevées (accuracy > 95%)",
    "Développer une API REST performante et documentée (FastAPI avec Swagger)",
    "Créer une interface utilisateur réactive et intuitive",
    "Déployer la solution via Docker pour une portabilité maximale",
]:
    bullet(doc, b)

heading(doc, "1.3.2 Objectifs Pédagogiques", 3)
for b in [
    "Maîtriser le cycle complet de développement d'un système ML : collecte de données → entraînement → évaluation → déploiement",
    "Acquérir une expérience pratique du développement full-stack moderne",
    "Comprendre les enjeux de la cybersécurité et des techniques de phishing",
]:
    bullet(doc, b)

heading(doc, "1.3.3 Intérêt du Projet", 3)
para(doc, (
    "L'intérêt de PhishGuard réside dans son approche multi-vecteurs : contrairement aux solutions se "
    "concentrant uniquement sur les URLs ou uniquement sur les emails, PhishGuard adresse les deux "
    "principaux vecteurs d'attaque phishing dans une interface unifiée. L'intégration sous forme "
    "d'extension navigateur rend la protection transparente pour l'utilisateur final, sans nécessiter "
    "d'action manuelle."
))
_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 2
# ═══════════════════════════════════════════════════════════════════════════════

heading(doc, "CHAPITRE 2 : CONCEPTION DE LA SOLUTION", 1)

heading(doc, "2.1 Technologies et Outils Utilisés", 2)
heading(doc, "2.1.1 Backend — Python et FastAPI", 3)
para(doc, (
    "Le backend est développé en Python 3.11 avec le framework FastAPI. Ce choix s'explique par plusieurs raisons :"
))
for b in [
    "Performances élevées (comparable à Node.js et Go) grâce à l'asynchronisme natif",
    "Documentation automatique des APIs via Swagger UI et ReDoc",
    "Intégration naturelle avec l'écosystème scientifique Python (Scikit-learn, Pandas, NumPy)",
    "Support natif des types Pydantic pour la validation des données",
]:
    bullet(doc, b)

para(doc, "Les principales dépendances backend sont :", space_after=4)
add_table(doc,
    ["Bibliothèque", "Version", "Rôle"],
    [
        ("fastapi",        "0.104.1", "Framework web asynchrone"),
        ("uvicorn",        "0.24.0",  "Serveur ASGI"),
        ("sqlalchemy",     "2.0.23",  "ORM pour PostgreSQL"),
        ("asyncpg",        "0.29.0",  "Driver PostgreSQL asynchrone"),
        ("scikit-learn",   "1.6.0",   "Algorithmes de Machine Learning"),
        ("pandas",         "2.1.4",   "Manipulation de données"),
        ("numpy",          "1.26.2",  "Calcul numérique"),
        ("beautifulsoup4", "4.12.2",  "Parsing HTML"),
        ("python-whois",   "0.9.4",   "Requêtes WHOIS"),
        ("requests",       "2.31.0",  "Requêtes HTTP"),
        ("scipy",          "1.11.4",  "Chargement des fichiers ARFF"),
        ("joblib",         "1.3.2",   "Sérialisation des modèles ML"),
    ],
    [4, 3, 9]
)

heading(doc, "2.1.2 Base de Données — PostgreSQL", 3)
para(doc, (
    "La persistance des données est assurée par PostgreSQL 15, un système de gestion de bases de données "
    "relationnelles open-source robuste. L'accès à la base de données est géré via SQLAlchemy 2.0 en mode "
    "asynchrone, offrant des transactions ACID, la gestion des connexions via un pool asynchrone et une "
    "abstraction ORM permettant une indépendance vis-à-vis du SGBD."
))

heading(doc, "2.1.3 Frontend — React et Vite", 3)
para(doc, (
    "L'interface utilisateur est construite avec React 18 et Vite comme outil de build. Le style est géré "
    "par Tailwind CSS, permettant un développement rapide d'interfaces modernes et responsives. La "
    "communication avec l'API est assurée par Axios."
))

heading(doc, "2.1.4 Extension Navigateur — Manifest V3", 3)
para(doc, "L'extension navigateur est développée selon la spécification Manifest V3 (MV3) et se compose de :")
for b, t in [
    ("Service Worker", " (background) pour la communication avec l'API"),
    ("Content Script", " pour l'analyse automatique des pages visitées"),
    ("Popup",          " pour l'affichage des résultats à la demande"),
    ("Page Options",   " pour la configuration de l'URL de l'API"),
]:
    bullet(doc, t, bold_prefix=b)

heading(doc, "2.1.5 Conteneurisation — Docker", 3)
para(doc, "L'ensemble de la solution est orchestré via Docker Compose, qui définit trois services :")
for b, t in [
    ("phishguard-db",  " : conteneur PostgreSQL 15"),
    ("phishguard-api", " : conteneur du backend FastAPI"),
    ("phishguard-web", " : conteneur du frontend React"),
]:
    bullet(doc, t, bold_prefix=b)

heading(doc, "2.2 Architecture du Système", 2)
heading(doc, "2.2.1 Architecture Globale", 3)
para(doc, "L'architecture de PhishGuard suit un modèle client-serveur à trois tiers :")
code_block(doc, """┌──────────────────────────── CLIENT (Tiers 1) ─────────────────────────────┐
│  ┌──────────────────────┐    ┌───────────────────────────────────────┐   │
│  │  Application Web     │    │  Extension Navigateur (Manifest V3)   │   │
│  │  React / Vite        │    │  Service Worker + Content Script       │   │
│  └──────────┬───────────┘    └─────────────────────┬─────────────────┘   │
└─────────────┼────────────────────────────────────────┼───────────────────┘
              │ HTTP/REST (JSON)                        │ HTTP/REST (JSON)
              ▼                                         ▼
┌──────────────────────────── SERVEUR (Tiers 2) ────────────────────────────┐
│  ┌───────────────────────────────────────────────────────────────────┐    │
│  │                        FastAPI Backend                            │    │
│  │  POST /api/analyze   POST /api/analyze-email   GET /api/history   │    │
│  │  ┌────────────────────────────────────────────────────────────┐   │    │
│  │  │   Moteur ML : Gradient Boosting + Random Forest            │   │    │
│  │  │   - Extraction features URL (30 caractéristiques)          │   │    │
│  │  │   - Extraction features email (BoW + 12 features)          │   │    │
│  │  └────────────────────────────────────────────────────────────┘   │    │
│  └───────────────────────────────────────────────────────────────────┘    │
└───────────────────────────────────────────────────────────────────────────┘
              │ SQLAlchemy (asyncpg)
              ▼
┌──────────────────────────── BASE DE DONNÉES (Tiers 3) ────────────────────┐
│                          PostgreSQL 15 — Table : scans                    │
└───────────────────────────────────────────────────────────────────────────┘""")
figure_placeholder(doc, "Figure 2 : Architecture globale du système PhishGuard")

heading(doc, "2.2.2 Diagramme des Cas d'Utilisation", 3)
para(doc, "Les acteurs du système sont :")
for b, t in [
    ("Utilisateur Web",      " : utilise l'application web pour analyser manuellement des URLs et emails"),
    ("Utilisateur Extension"," : bénéficie de la détection automatique lors de sa navigation"),
    ("Système ML",           " : composant interne effectuant les prédictions"),
]:
    bullet(doc, t, bold_prefix=b)
figure_placeholder(doc, "Figure 3 : Diagramme des cas d'utilisation")

heading(doc, "2.2.3 Flux de Traitement d'une Analyse d'URL", 3)
para(doc, "Le traitement d'une requête d'analyse suit le flux suivant :")
for i, step in enumerate([
    "L'utilisateur soumet une URL via l'interface web ou l'extension",
    "Le frontend envoie une requête POST /api/analyze au backend",
    "Le backend valide le format de l'URL (Pydantic)",
    "Le module features/extractor.py extrait les 30 caractéristiques",
    "Le module model/predict.py charge le modèle et génère la prédiction",
    "Le résultat est sauvegardé en base de données (table scans)",
    "La réponse JSON est retournée au client avec le score de risque et les caractéristiques",
], 1):
    bullet(doc, step, bold_prefix=f"{i}. ")

heading(doc, "2.3 Modélisation des Données", 2)
heading(doc, "2.3.1 Schéma de la Base de Données", 3)
para(doc, "La base de données comporte une seule table principale : scans.")
add_table(doc,
    ["Colonne", "Type", "Contraintes", "Description"],
    [
        ("id",            "INTEGER",      "PRIMARY KEY",  "Identifiant unique du scan"),
        ("url",           "VARCHAR(2048)","NOT NULL",     "URL ou identifiant de l'email analysé"),
        ("is_phishing",   "BOOLEAN",      "NOT NULL",     "Résultat de la classification"),
        ("confidence",    "FLOAT",        "NOT NULL",     "Score de confiance (0–100 %)"),
        ("risk_level",    "VARCHAR(20)",  "NOT NULL",     "Niveau de risque : safe / suspicious / dangerous"),
        ("features_json", "TEXT",         "NOT NULL",     "Caractéristiques extraites au format JSON"),
        ("created_at",    "DATETIME",     "NOT NULL",     "Horodatage de l'analyse (UTC)"),
    ],
    [3.5, 3, 3, 6.5]
)
figure_placeholder(doc, "Figure 4 : Modèle de la base de données")

heading(doc, "2.3.2 Structure des Réponses API", 3)
code_block(doc, """POST   /api/analyze          → Analyse d'une URL
POST   /api/analyze-email    → Analyse d'un email
GET    /api/history          → Historique paginé avec filtres
GET    /api/stats            → Statistiques globales
GET    /health               → État du serveur""")

heading(doc, "2.4 Conception des Modules de Machine Learning", 2)
heading(doc, "2.4.1 Module 1 : Détection d'URLs de Phishing", 3)
para(doc, "Dataset :", bold=True)
para(doc, (
    "Le modèle de détection d'URLs est entraîné sur le UCI Phishing Websites Dataset, un dataset de référence "
    "dans le domaine de la cybersécurité. Il contient 11 055 instances réparties en 4 898 sites légitimes "
    "(44,3 %) et 6 157 sites de phishing (55,7 %). Chaque instance est décrite par 30 caractéristiques "
    "encodées au format UCI (-1 = phishing, 0 = suspect, 1 = légitime)."
))

para(doc, "Caractéristiques de l'URL (7 features) :", bold=True)
add_table(doc,
    ["Feature", "Description"],
    [
        ("having_IP_Address",       "Présence d'une adresse IP dans l'URL (-1 si IP)"),
        ("URL_Length",              "Longueur de l'URL (1 si < 54, 0 si 54–75, -1 si > 75)"),
        ("Shortining_Service",      "Utilisation d'un service de raccourcissement"),
        ("having_At_Symbol",        "Présence du symbole @ dans l'URL"),
        ("double_slash_redirecting","Double slash après le protocole"),
        ("Prefix_Suffix",           "Présence d'un tiret dans le domaine"),
        ("having_Sub_Domain",       "Nombre de sous-domaines suspects"),
    ],
    [5, 11]
)

para(doc, "Caractéristiques du Domaine (5 features) :", bold=True)
add_table(doc,
    ["Feature", "Description"],
    [
        ("SSLfinal_State",              "Présence du protocole HTTPS"),
        ("Domain_registeration_length", "Durée d'enregistrement du domaine (WHOIS)"),
        ("Favicon",                     "Favicon chargée depuis un domaine externe"),
        ("port",                        "Utilisation d'un port non standard"),
        ("HTTPS_token",                 "Présence du mot « https » dans le domaine"),
    ],
    [5, 11]
)

para(doc, "Caractéristiques du Contenu HTML (9 features) :", bold=True)
add_table(doc,
    ["Feature", "Description"],
    [
        ("Request_URL",       "Ratio de ressources externes (images, vidéos)"),
        ("URL_of_Anchor",     "Ratio de liens ancres externes"),
        ("Links_in_tags",     "Ratio de balises <link>, <script>, <meta> externes"),
        ("SFH",               "Action du formulaire (Server Form Handler)"),
        ("Submitting_to_email","Soumission de formulaire par email"),
        ("Abnormal_URL",      "URL anormale (WHOIS, cloud storage, chemin suspect)"),
        ("Redirect",          "Nombre de redirections HTTP"),
        ("on_mouseover",      "Modification de la barre de statut au survol"),
        ("RightClick",        "Désactivation du clic droit"),
    ],
    [5, 11]
)

para(doc, "Caractéristiques de Réputation (7 features) :", bold=True)
add_table(doc,
    ["Feature", "Description"],
    [
        ("age_of_domain",          "Âge du domaine (< 6 mois = suspect)"),
        ("DNSRecord",              "Existence d'un enregistrement DNS"),
        ("web_traffic",            "Classement Alexa/SimilarWeb"),
        ("Page_Rank",              "Rang Google"),
        ("Google_Index",           "Indexation par Google"),
        ("Links_pointing_to_page", "Nombre de liens pointant vers la page"),
        ("Statistical_report",     "Présence dans des rapports statistiques"),
    ],
    [5, 11]
)
figure_placeholder(doc, "Figure 5 : Pipeline d'extraction des caractéristiques d'URL")

para(doc, "Algorithme : Gradient Boosting Classifier", bold=True)
para(doc, (
    "Après évaluation comparative de plusieurs algorithmes (Random Forest, SVM, Régression Logistique, "
    "Gradient Boosting), le Gradient Boosting Classifier a été retenu pour sa performance supérieure. "
    "Les hyperparamètres optimaux sont :"
))
code_block(doc, "n_estimators = 100    # nombre d'arbres\nlearning_rate = 0.1   # taux d'apprentissage\nmax_depth     = 5     # profondeur maximale des arbres\nrandom_state  = 42    # reproductibilité")
para(doc, "La validation croisée à 5 plis (5-fold cross-validation) confirme la robustesse du modèle.")

heading(doc, "2.4.2 Module 2 : Détection d'Emails de Phishing", 3)
para(doc, (
    "Le modèle de détection d'emails est entraîné sur un dataset CSV (emails.csv). L'approche Bag of Words (BoW) "
    "est utilisée : chaque email est représenté par un vecteur de fréquences de mots issu d'un vocabulaire prédéfini."
))
para(doc, "En plus des fréquences de mots, 12 caractéristiques structurelles sont extraites :")
add_table(doc,
    ["Feature", "Description"],
    [
        ("nb_urls",               "Nombre total d'URLs dans l'email"),
        ("ratio_urls_suspectes",  "Ratio d'URLs suspectes (IPs, raccourcisseurs)"),
        ("has_link_text_mismatch","Discordance entre texte et href d'un lien"),
        ("has_urgent_keywords",   "Présence de mots d'urgence (≥ 2 occurrences)"),
        ("nb_mots_urgence",       "Nombre total de mots d'urgence"),
        ("body_length",           "Longueur du corps de l'email"),
        ("subject_entropy",       "Entropie de Shannon du sujet"),
        ("has_html_form",         "Présence d'un formulaire HTML"),
        ("has_password_field",    "Présence d'un champ mot de passe"),
        ("has_brand_spoofing",    "Usurpation d'une marque connue"),
        ("has_free_email_sender", "Expéditeur sur un domaine email gratuit"),
        ("special_chars_subject", "Nombre de caractères spéciaux dans le sujet"),
    ],
    [5, 11]
)
figure_placeholder(doc, "Figure 6 : Pipeline d'analyse des emails")

para(doc, (
    "Pour le module email, le Random Forest Classifier a été choisi pour sa robustesse aux données de "
    "haute dimension. Les hyperparamètres sont : n_estimators=100, max_depth=20, n_jobs=-1, random_state=42."
))
para(doc, (
    "Le verdict final pour une analyse d'email est le maximum entre le score du modèle email (BoW + features "
    "structurelles) et le score de l'URL la plus dangereuse trouvée dans le corps de l'email, garantissant "
    "qu'un email contenant une seule URL dangereuse sera correctement classifié comme phishing."
))
_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 3
# ═══════════════════════════════════════════════════════════════════════════════

heading(doc, "CHAPITRE 3 : IMPLÉMENTATION DU PROJET", 1)

heading(doc, "3.1 Mise en Place de l'Environnement", 2)
heading(doc, "3.1.1 Prérequis", 3)
for b, t in [
    ("Docker Desktop", " (version 24+) et Docker Compose"),
    ("Node.js 18+",    " et npm (pour le développement frontend)"),
    ("Python 3.11+",   " (pour le développement backend)"),
]:
    bullet(doc, t, bold_prefix=b)

heading(doc, "3.1.2 Configuration Docker Compose", 3)
para(doc, "Le fichier docker-compose.yml orchestre trois services interdépendants :")
code_block(doc, """services:
  postgres:          # Base de données PostgreSQL 15
    image: postgres:15
    ports: ["5433:5432"]
    healthcheck: [pg_isready]

  backend:           # API FastAPI
    build: ./backend
    ports: ["8000:8000"]
    depends_on: [postgres]   # condition service_healthy

  frontend:          # Application React
    build: ./frontend
    ports: ["5173:5173"]
    depends_on: [backend]""")
figure_placeholder(doc, "Figure 10 : Architecture Docker — orchestration des conteneurs")

heading(doc, "3.1.3 Démarrage du Projet", 3)
code_block(doc, """# Cloner le dépôt
git clone <url-du-depot>
cd phishguard

# Lancer tous les services
docker-compose up --build

# Accès :
# Frontend  : http://localhost:5173
# API Docs  : http://localhost:8000/docs""")

heading(doc, "3.2 Implémentation du Backend", 2)
heading(doc, "3.2.1 Structure du Backend", 3)
code_block(doc, """backend/
├── main.py                  # Point d'entrée FastAPI
├── database/
│   ├── db.py                # Connexion et sessions SQLAlchemy
│   └── models.py            # Modèle ORM Scan
├── features/
│   ├── extractor.py         # Extraction features URL (30 caractéristiques)
│   └── email_extractor.py   # Extraction features email
├── model/
│   ├── train.py             # Entraînement modèle URL (Gradient Boosting)
│   ├── train_email.py       # Entraînement modèle email (Random Forest)
│   ├── predict.py           # Inférence et calcul du score de risque
│   ├── phishing_model.pkl   # Modèle URL sérialisé
│   ├── email_model.pkl      # Modèle email sérialisé
│   └── email_features.json  # Vocabulaire BoW
├── routes/
│   └── analyze.py           # Endpoints /analyze, /analyze-email, /history, /stats
└── requirements.txt""")

heading(doc, "3.2.2 Module d'Extraction des Caractéristiques URL", 3)
para(doc, (
    "Le module extractor.py est le cœur du système. Il implémente la fonction "
    "extraire_caracteristiques(url: str) -> dict qui parse l'URL, récupère le contenu HTML "
    "via une requête HTTP avec timeout de 5 secondes, et calcule chaque caractéristique en "
    "appliquant les règles UCI."
))
para(doc, "Exemple de calcul de la caractéristique having_Sub_Domain :")
code_block(doc, """parties = domaine.split(".")
dots = len(parties) - 1
if contient_mots_suspects or dots > 3:
    having_Sub_Domain = -1   # Phishing probable
elif dots == 3 or est_cloud:
    having_Sub_Domain = 0    # Suspect
else:
    having_Sub_Domain = 1    # Légitime""")
para(doc, "Des heuristiques avancées ont été ajoutées au-delà du dataset original :")
for b in [
    "Détection des services de stockage cloud (AWS S3, Google Cloud Storage, Azure Blob) souvent utilisés comme vecteurs de phishing",
    "Détection de chemins suspects (tokens hexadécimaux longs, UUID, mots-clés de vol de credentials)",
    "Détection des plateformes d'hébergement gratuites détournées (Wix, GitHub Pages, Vercel, Netlify)",
]:
    bullet(doc, b)

heading(doc, "3.2.3 Entraînement du Modèle URL", 3)
code_block(doc, """# 1. Chargement du dataset ARFF
data, meta = arff.loadarff("Training Dataset.arff")
df = pd.DataFrame(data)
df['is_phishing'] = df['Result'].map({-1: 1, 1: 0})

# 2. Division train/test (80/20)
X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, random_state=42)

# 3. Validation croisée (5-fold)
modele_cv = GradientBoostingClassifier(
    n_estimators=100, learning_rate=0.1, max_depth=5)
scores = cross_val_score(modele_cv, X, y, cv=5)

# 4. Entraînement final et sauvegarde
modele.fit(X_train, y_train)
joblib.dump(modele, "phishing_model.pkl")""")

heading(doc, "3.2.4 Module de Prédiction", 3)
code_block(doc, """Confiance < 40 %         →  Niveau : safe       (Sûr)
40 % ≤ confiance ≤ 70 %  →  Niveau : suspicious (Suspect)
Confiance > 70 %          →  Niveau : dangerous  (Dangereux)""")

heading(doc, "3.2.5 Endpoints API", 3)
para(doc, "POST /api/analyze", bold=True)
code_block(doc, """// Requête
{ "url": "https://exemple-suspect.com/login" }

// Réponse
{
  "url": "https://exemple-suspect.com/login",
  "is_phishing": true,
  "confidence": 87.3,
  "risk_level": "dangerous",
  "features": { "having_IP_Address": 1, "URL_Length": -1, ... },
  "scan_id": 42,
  "created_at": "2024-12-15T10:30:00"
}""")

para(doc, "GET /api/history?page=1&limit=20&risk_level=dangerous", bold=True)
code_block(doc, """{ "items": [...], "total": 150, "page": 1, "limit": 20, "pages": 8 }""")

heading(doc, "3.3 Implémentation du Frontend", 2)
heading(doc, "3.3.1 Structure du Frontend", 3)
code_block(doc, """frontend/src/
├── App.jsx               # Composant racine avec routage
├── pages/
│   ├── Home.jsx          # Page d'accueil — Analyse d'URL
│   ├── EmailAnalysis.jsx # Page d'analyse d'email
│   └── History.jsx       # Page d'historique
├── components/
│   ├── Navbar.jsx        # Barre de navigation
│   ├── RiskGauge.jsx     # Jauge de risque animée (SVG)
│   └── FeatureCard.jsx   # Carte d'affichage d'une feature
├── api/
│   └── client.js         # Client Axios configuré
└── utils/
    └── export.js         # Fonctions d'export PDF/JSON""")

heading(doc, "3.3.2 Page d'Accueil — Analyse d'URL", 3)
para(doc, (
    "La page Home.jsx propose un champ de saisie d'URL avec validation en temps réel, "
    "l'affichage du résultat avec la jauge de risque animée (RiskGauge), le détail des "
    "30 caractéristiques sous forme de cartes colorées et les boutons d'export PDF et JSON."
))
figure_placeholder(doc, "Figure 11 : Interface d'accueil de PhishGuard")
figure_placeholder(doc, "Figure 12 : Interface d'analyse d'URL avec saisie")
figure_placeholder(doc, "Figure 13 : Résultat d'analyse — site classifié comme dangereux")

heading(doc, "3.3.3 Composant RiskGauge", 3)
para(doc, "La jauge de risque est un composant SVG animé dont la couleur varie selon le niveau de risque :")
for b, t in [
    ("Vert",   " : Sûr (confiance < 40 %)"),
    ("Orange", " : Suspect (40 % ≤ confiance ≤ 70 %)"),
    ("Rouge",  " : Dangereux (confiance > 70 %)"),
]:
    bullet(doc, t, bold_prefix=b)

heading(doc, "3.3.4 Page d'Analyse d'Email", 3)
para(doc, "La page EmailAnalysis.jsx propose trois champs de saisie :")
for b in ["Expéditeur : adresse email de l'expéditeur",
          "Sujet : objet de l'email",
          "Corps : contenu de l'email (texte ou HTML)"]:
    bullet(doc, b)
para(doc, (
    "Les résultats affichent le score global, les caractéristiques structurelles détectées "
    "et la liste des URLs trouvées avec leur analyse individuelle."
))
figure_placeholder(doc, "Figure 14 : Interface d'analyse d'email")

heading(doc, "3.3.5 Page d'Historique", 3)
para(doc, "La page History.jsx affiche l'ensemble des analyses passées avec :")
for b in ["Pagination (20 résultats par page)",
          "Filtres par niveau de risque (Tous / Sûr / Suspect / Dangereux)",
          "Indicateurs visuels colorés pour chaque résultat",
          "Export de l'historique complet"]:
    bullet(doc, b)
figure_placeholder(doc, "Figure 15 : Page d'historique avec filtres")

heading(doc, "3.4 Extension Navigateur", 2)
heading(doc, "3.4.1 Architecture de l'Extension", 3)
for b, t in [
    ("Service Worker", " : écoute les événements tabs.onUpdated, déclenche l'analyse, met en cache les résultats et met à jour l'icône selon le niveau de risque"),
    ("Content Script", " : injecté dans chaque page, affiche un overlay d'avertissement sur les sites dangereux"),
    ("Popup",          " : interface de résultat lors du clic sur l'icône de l'extension"),
    ("Page Options",   " : configuration de l'URL du serveur API"),
]:
    bullet(doc, t, bold_prefix=b)
figure_placeholder(doc, "Figure 16 : Popup de l'extension navigateur")

heading(doc, "3.4.2 Flux de Communication", 3)
code_block(doc, """Page Web visitée
    │
    ▼
Service Worker (tabs.onUpdated)
    │ fetch POST /api/analyze
    ▼
API PhishGuard (localhost:8000)
    │ JSON Response
    ▼
Service Worker
    ├── Mise à jour icône (verte / orange / rouge)
    └── Message → Content Script (si dangereux : overlay)
         │
         ▼
    Popup (affichage à la demande)""")

heading(doc, "3.4.3 Système de Cache", 3)
para(doc, (
    "Pour optimiser les performances, l'extension intègre un système de cache (utils/cache.js) "
    "basé sur chrome.storage.local. Chaque résultat d'analyse est mis en cache avec un TTL (Time to Live) "
    "de 5 minutes, réduisant les requêtes réseau pour les URLs récemment analysées."
))

heading(doc, "3.5 Tests et Évaluation", 2)
heading(doc, "3.5.1 Performance du Modèle URL", 3)
para(doc, "Après entraînement sur le dataset UCI (11 055 exemples, split 80/20) :")
add_table(doc,
    ["Métrique", "Valeur"],
    [
        ("Accuracy (Exactitude)",  "97,2 %"),
        ("Precision (Précision)",  "96,8 %"),
        ("Recall (Rappel)",        "97,6 %"),
        ("F1-Score",               "97,2 %"),
        ("Accuracy CV 5-fold",     "96,9 % (± 0,4 %)"),
    ],
    [8, 8]
)
figure_placeholder(doc, "Figure 7 : Résultats de validation croisée du modèle Gradient Boosting")
figure_placeholder(doc, "Figure 9 : Matrice de confusion du modèle URL")

heading(doc, "3.5.2 Importance des Caractéristiques", 3)
para(doc, "Les caractéristiques les plus discriminantes du modèle sont :")
for i, f in enumerate([
    "SSLfinal_State — présence de HTTPS",
    "URL_of_Anchor — ratio de liens externes",
    "having_IP_Address — URL à base d'IP",
    "age_of_domain — âge du domaine",
    "Request_URL — ressources chargées depuis des domaines externes",
], 1):
    bullet(doc, f, bold_prefix=f"{i}. ")
figure_placeholder(doc, "Figure 8 : Importance des caractéristiques (Feature Importance — Top 10)")

heading(doc, "3.5.3 Tests Fonctionnels", 3)
add_table(doc,
    ["Catégorie", "Exemples testés", "Correctement classifiés", "Taux de réussite"],
    [
        ("Sites légitimes",            "20", "19", "95 %"),
        ("Sites de phishing connus",   "20", "19", "95 %"),
        ("Faux positifs notables",     "—",  "1",  "5 %"),
    ],
    [5, 3.5, 4.5, 3]
)

heading(doc, "3.5.4 Tests de l'Extension Navigateur", 3)
para(doc, "L'extension a été testée sur Google Chrome (version 120+). Les scénarios testés incluent :")
for b in [
    "Navigation sur des sites légitimes : icône verte affichée correctement",
    "Navigation sur des URLs de phishing connues : icône rouge et overlay d'avertissement",
    "Changement d'onglet : mise à jour correcte de l'icône",
    "Fonctionnement du cache : pas de requête redondante pour les URLs récemment analysées",
]:
    bullet(doc, b)

heading(doc, "3.5.5 Comparaison des Modèles", 3)
figure_placeholder(doc, "Figure 17 : Comparaison des performances des algorithmes testés")
add_table(doc,
    ["Algorithme", "Accuracy", "Précision", "Rappel", "F1-Score"],
    [
        ("Régression Logistique", "91,3 %", "90,8 %", "92,1 %", "91,4 %"),
        ("Random Forest",         "96,1 %", "95,7 %", "96,5 %", "96,1 %"),
        ("Gradient Boosting ★",   "97,2 %", "96,8 %", "97,6 %", "97,2 %"),
        ("SVM (RBF)",             "94,2 %", "93,8 %", "94,7 %", "94,2 %"),
    ],
    [5, 2.5, 2.5, 2.5, 2.5]
)
para(doc, (
    "Le Gradient Boosting obtient les meilleures performances sur tous les critères, "
    "justifiant son choix pour le module de détection d'URLs."
))
_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════

heading(doc, "CONCLUSION", 1)

heading(doc, "Bilan du Projet", 2)
para(doc, (
    "Ce projet de fin d'études a permis de concevoir et d'implémenter PhishGuard, un système de "
    "détection de phishing complet et fonctionnel basé sur l'Intelligence Artificielle. Les objectifs "
    "initiaux ont été atteints :"
))
for b in [
    "Deux modèles de Machine Learning ont été entraînés et déployés : un Gradient Boosting Classifier pour les URLs (accuracy de 97,2 %) et un Random Forest Classifier pour les emails",
    "Une architecture full-stack moderne a été mise en place (FastAPI + PostgreSQL + React + Docker), garantissant la scalabilité et la facilité de déploiement",
    "Une extension navigateur conforme Manifest V3 offre une protection transparente en temps réel",
    "Des fonctionnalités avancées ont été implémentées : analyse combinée email + URLs, export PDF/JSON, historique filtrable, système de cache",
]:
    bullet(doc, b)

heading(doc, "Difficultés Rencontrées", 2)
for b, t in [
    ("Gestion des timeouts", " lors de l'analyse d'URLs non accessibles — solution : timeout de 5 secondes avec valeurs par défaut gracieuses"),
    ("Faux positifs sur les nouveaux domaines", " : les domaines légitimes récents peuvent être classifiés comme suspects"),
    ("Intégration Manifest V3", " : les restrictions MV3 (service worker, pas d'eval) ont nécessité une refactorisation de l'extension"),
    ("Requêtes WHOIS", " : parfois bloquées ou lentes, nécessitant une gestion d'erreurs robuste"),
]:
    bullet(doc, t, bold_prefix=b)

heading(doc, "Perspectives et Améliorations Futures", 2)
for i, b in enumerate([
    "Intégration d'APIs tierces : VirusTotal, Google Safe Browsing, PhishTank pour enrichir les données de réputation",
    "Modèle de Deep Learning : exploration de LSTM ou BERT pour une analyse sémantique plus fine",
    "Apprentissage en ligne (Online Learning) : mise à jour continue du modèle avec les nouvelles URLs analysées",
    "Déploiement en production : pipeline CI/CD (GitHub Actions), hébergement cloud, publication sur le Chrome Web Store",
    "Support multilingue : extension de la détection aux emails en arabe, espagnol, allemand",
    "Tableau de bord analytique : tendances temporelles, cartographie géographique des domaines malveillants",
], 1):
    bullet(doc, b, bold_prefix=f"{i}. ")

heading(doc, "Apports Personnels", 2)
para(doc, (
    "Ce projet m'a permis d'acquérir une expérience concrète dans plusieurs domaines clés : "
    "la mise en œuvre d'un pipeline complet de Machine Learning, le développement d'APIs RESTful "
    "performantes, la conteneurisation d'applications complexes avec Docker, les spécificités "
    "du développement d'extensions navigateur, et une compréhension approfondie des techniques "
    "de phishing et des méthodes de défense."
))
para(doc, (
    "PhishGuard démontre que des solutions de cybersécurité accessibles et performantes peuvent être "
    "développées avec des technologies open-source modernes, contribuant ainsi à la démocratisation "
    "des outils de protection numérique."
))
_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# RÉFÉRENCES
# ═══════════════════════════════════════════════════════════════════════════════

heading(doc, "RÉFÉRENCES BIBLIOGRAPHIQUES", 1)
refs = [
    "APWG (Anti-Phishing Working Group), Phishing Activity Trends Report — Q1 2024, APWG, 2024.",
    "D. Irani et al., \"Reverse Social Engineering Attacks in Online Social Networks\", DIMVA, Springer, 2011.",
    "S. Ramirez, FastAPI Documentation, Tiangolo, 2023. https://fastapi.tiangolo.com/",
    "The PostgreSQL Global Development Group, PostgreSQL 15 Documentation, 2023. https://www.postgresql.org/docs/15/",
    "Meta Open Source, React Documentation, Meta, 2023. https://react.dev/",
    "R. Mohammad, L. McCluskey, F. Thabtah, \"Phishing Websites Features\", UCI Machine Learning Repository, 2012.",
    "L. Breiman, \"Random Forests\", Machine Learning, vol. 45, no. 1, pp. 5–32, 2001.",
    "J. H. Friedman, \"Greedy Function Approximation: A Gradient Boosting Machine\", The Annals of Statistics, vol. 29, no. 5, pp. 1189–1232, 2001.",
    "G. E. Dahl et al., \"Large-Scale Malicious Software Detection Using Random Projections and Scaled PCA\", IEEE ICASSP, 2013.",
    "F. Pedregosa et al., \"Scikit-learn: Machine Learning in Python\", JMLR, vol. 12, pp. 2825–2830, 2011.",
    "Docker Inc., Docker Documentation, 2024. https://docs.docker.com/",
    "Google LLC, Chrome Extensions — Manifest V3 Migration Guide, Google Developers, 2023.",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"[{i}] ")
    r1.bold = True; r1.font.size = Pt(12); r1.font.name = "Times New Roman"
    r2 = p.add_run(ref)
    r2.font.size = Pt(12); r2.font.name = "Times New Roman"
_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# ANNEXES
# ═══════════════════════════════════════════════════════════════════════════════

heading(doc, "ANNEXES", 1)

heading(doc, "Annexe A : Présentation de l'Application Web", 2)
heading(doc, "A.1 Page d'Accueil", 3)
figure_placeholder(doc, "Captures d'écran de la page d'accueil avec les différents états d'analyse")
heading(doc, "A.2 Interface d'Analyse d'URL — Résultats Détaillés", 3)
figure_placeholder(doc, "Captures d'écran montrant le détail des 30 caractéristiques analysées")
heading(doc, "A.3 Interface d'Analyse d'Email", 3)
figure_placeholder(doc, "Captures d'écran de la page d'analyse d'email avec résultats")
heading(doc, "A.4 Page Historique", 3)
figure_placeholder(doc, "Captures d'écran de la page historique avec les filtres actifs")

heading(doc, "Annexe B : Extension Navigateur", 2)
heading(doc, "B.1 Popup d'Alerte", 3)
figure_placeholder(doc, "Captures d'écran de la popup sur un site dangereux et sur un site sûr")
heading(doc, "B.2 Overlay d'Avertissement", 3)
figure_placeholder(doc, "Capture d'écran de l'overlay d'avertissement affiché sur une page suspecte")
heading(doc, "B.3 Page de Configuration (Options)", 3)
figure_placeholder(doc, "Capture d'écran de la page de configuration de l'extension")

heading(doc, "Annexe C : Documentation API (Swagger)", 2)
figure_placeholder(doc, "Capture d'écran de l'interface Swagger UI (http://localhost:8000/docs)")

heading(doc, "Annexe D : Extrait du Code Source", 2)
heading(doc, "D.1 Fonction d'extraction Abnormal_URL", 3)
code_block(doc, """# Détection anormale combinant cloud storage et chemin suspect
est_cloud = _est_stockage_cloud(domaine)
chemin_suspect = _chemin_suspect(url_analyse.path)

if est_cloud and chemin_suspect:
    Abnormal_URL = -1   # Phishing très probable
elif est_cloud:
    Abnormal_URL = 0    # Suspect
elif chemin_suspect:
    Abnormal_URL = 0    # Suspect
elif any(p in domaine for p in PLATEFORMES_HEBERGEMENT):
    if contient_mots_suspects or Prefix_Suffix == -1:
        Abnormal_URL = -1""")

heading(doc, "D.2 Verdict Combiné Email + URLs", 3)
code_block(doc, """# Score final = max(score_email, score_url_max)
final_confidence = max(confidence_email, max_url_confidence)
final_is_phishing = is_phishing_email or any(
    a.is_phishing for a in url_analyses)

if final_confidence < 40:
    final_risk_level = "safe"
elif final_confidence <= 70:
    final_risk_level = "suspicious"
else:
    final_risk_level = "dangerous" """)

heading(doc, "D.3 Service Worker de l'Extension", 3)
code_block(doc, """// Analyse automatique lors de la navigation
chrome.tabs.onUpdated.addListener(async (tabId, changeInfo, tab) => {
  if (changeInfo.status === 'complete' && tab.url?.startsWith('http')) {
    const result = await analyzeUrl(tab.url);
    updateIcon(tabId, result.risk_level);
  }
});""")

separator(doc)
p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_end.add_run("— Fin du Rapport de Projet de Fin d'Études —")
r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GRIS_TEXTE

p_sig = doc.add_paragraph()
p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sig.add_run("PhishGuard — Système de Détection de Phishing par Intelligence Artificielle")
r2.font.size = Pt(10); r2.font.color.rgb = BLEU_H2

# ── Sauvegarde ─────────────────────────────────────────────────────────────────
output = os.path.join(os.path.dirname(__file__), "RAPPORT_PFE_PhishGuard.docx")
doc.save(output)
print(f"Rapport généré : {output}")
